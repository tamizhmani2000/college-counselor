"""
generate_plan.py — College Planning Document Generator

Produces a formatted Word document (.docx) for a student's college plan.

Usage
-----
Populate the `generate_plan()` call at the bottom of this file with the
student's data, then run:

    pip install python-docx
    python tools/generate_plan.py

Output is written to:  output/<StudentName>_plan.docx

Data contracts
--------------
career_paths    : list[dict]  — keys: title, fit_label, major, grad_school,
                                      salary_start, salary_mid, outlook,
                                      why_fit, day_in_life
reach/match/safety_colleges : list[dict]  — keys: name, location, type,
                                      acceptance_rate, avg_gpa, sat_act_range,
                                      enrollment, program_notes, annual_coa,
                                      net_price, aid_notes, fit_reason
scholarships    : list[dict]  — keys: name, amount, eligibility, deadline, website
timeline_rows   : list[tuple] — (timeframe_label, action_items_string)
all other lists : list[str]
"""

import datetime
import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str):
    """Fill a table cell with a solid background colour (e.g. '1F3864')."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_col_width(table, col_index: int, width_cm: float):
    for row in table.rows:
        row.cells[col_index].width = Cm(width_cm)


def _add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3864')
    pb.append(bottom)
    pPr.append(pb)


# ---------------------------------------------------------------------------
# Document building blocks
# ---------------------------------------------------------------------------

def _add_cover_page(doc, student_name: str, grade: str,
                    city: str, state: str, intended_major: str):
    for _ in range(3):
        doc.add_paragraph()

    title = doc.add_paragraph('COLLEGE PLANNING GUIDE')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    doc.add_paragraph()

    name_p = doc.add_paragraph(student_name)
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name_p.runs[0]
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    meta = doc.add_paragraph(
        f'Grade {grade}  |  {city}, {state}  |  Intended Major: {intended_major}'
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(12)

    date_p = doc.add_paragraph(
        f'Prepared: {datetime.date.today().strftime("%B %d, %Y")}'
    )
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.size = Pt(11)
    date_p.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    doc.add_page_break()


def _add_section_heading(doc, number: int | str, title: str):
    """Level-1 heading with section number in navy."""
    p = doc.add_heading(f'Section {number}: {title}', level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    _add_horizontal_rule(doc)
    doc.add_paragraph()


def _add_sub_heading(doc, title: str):
    """Level-2 heading in blue."""
    p = doc.add_heading(title, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)


def _add_bullet(doc, text: str, level: int = 0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    doc.add_paragraph(text, style=style)


def _add_bold_label(doc, label: str, value: str):
    """One paragraph: bold label + normal value text."""
    p = doc.add_paragraph()
    r1 = p.add_run(f'{label}: ')
    r1.bold = True
    p.add_run(value)


def _add_career_path(doc, number: int, title: str, fit_label: str,
                     major: str, grad_school: str, salary_start: str,
                     salary_mid: str, outlook: str, why_fit: str,
                     day_in_life: str):
    _add_sub_heading(doc, f'{number}.  {title}  —  {fit_label}')
    _add_bold_label(doc, 'Recommended Major(s)', major)
    _add_bold_label(doc, 'Graduate School', grad_school)
    _add_bold_label(doc, 'Starting Salary', salary_start)
    _add_bold_label(doc, 'Mid-Career Salary', salary_mid)
    _add_bold_label(doc, 'Job Market Outlook', outlook)
    _add_bold_label(doc, 'Why This Fits You', why_fit)
    _add_bold_label(doc, 'Day in the Life', day_in_life)
    doc.add_paragraph()


def _add_college_table(doc, tier_label: str, tier_color_hex: str,
                       colleges: list[dict]):
    _add_sub_heading(doc, tier_label)

    if not colleges:
        doc.add_paragraph('No colleges in this tier.')
        return

    headers = [
        'College', 'Location', 'Accept %', 'Avg GPA',
        'SAT/ACT', 'Annual COA', 'Net Price', 'Why a Good Fit',
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_bg(cell, tier_color_hex)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)

    col_keys = [
        'name', 'location', 'acceptance_rate', 'avg_gpa',
        'sat_act_range', 'annual_coa', 'net_price', 'fit_reason',
    ]
    for college in colleges:
        row = table.add_row()
        for i, key in enumerate(col_keys):
            cell = row.cells[i]
            cell.text = college.get(key, '')
            cell.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph(
        '* Net price shown is estimated for the household income range provided.'
    )
    doc.add_paragraph()


def _add_scholarship_table(doc, scholarships: list[dict]):
    headers = ['Scholarship', 'Amount', 'Eligibility', 'Deadline', 'Where to Apply']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_bg(cell, '2E75B6')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)

    col_keys = ['name', 'amount', 'eligibility', 'deadline', 'website']
    for s in scholarships:
        row = table.add_row()
        for i, key in enumerate(col_keys):
            cell = row.cells[i]
            cell.text = s.get(key, '')
            cell.paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()


def _add_timeline_table(doc, timeline_rows: list[tuple]):
    headers = ['Timeframe', 'Action Items']
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_bg(cell, '1F3864')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    for timeframe, actions in timeline_rows:
        row = table.add_row()
        row.cells[0].text = timeframe
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].text = actions
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Main public function
# ---------------------------------------------------------------------------

def generate_plan(
    student_name: str,
    grade: str,
    city: str,
    state: str,
    intended_major: str,
    career_paths: list[dict],
    course_recs: list[str],
    gpa_targets: list[str],
    summer_programs: list[str],
    reach_colleges: list[dict],
    match_colleges: list[dict],
    safety_colleges: list[dict],
    activities: list[str],
    leadership_recs: list[str],
    service_recs: list[str],
    federal_aid_notes: list[str],
    state_aid_notes: list[str],
    institutional_aid: list[str],
    scholarships: list[dict],
    timeline_rows: list[tuple],
    strategy_notes: list[str],
) -> str:
    os.makedirs('output', exist_ok=True)
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    _add_cover_page(doc, student_name, grade, city, state, intended_major)

    _add_section_heading(doc, 1, 'Career Path Recommendations')
    doc.add_paragraph(
        f'Based on your academic strengths, interests, and goals, here are the three '
        f'career paths that are the best fit for you, {student_name.split()[0]}.'
    )
    doc.add_paragraph()
    for i, cp in enumerate(career_paths, 1):
        _add_career_path(doc, i, **cp)

    _add_section_heading(doc, 2, 'Academic Roadmap')

    _add_sub_heading(doc, '2.1  Recommended Courses')
    for item in course_recs:
        _add_bullet(doc, item)
    doc.add_paragraph()

    _add_sub_heading(doc, '2.2  GPA and Test Score Targets')
    for item in gpa_targets:
        _add_bullet(doc, item)
    doc.add_paragraph()

    _add_sub_heading(doc, '2.3  Summer Program and Internship Recommendations')
    for item in summer_programs:
        _add_bullet(doc, item)
    doc.add_paragraph()

    total = len(reach_colleges) + len(match_colleges) + len(safety_colleges)
    _add_section_heading(doc, 3, 'College List')
    doc.add_paragraph(
        f'The following {total} colleges have been selected based on your academic '
        f'profile, location preferences, intended major, and financial situation.'
    )
    doc.add_paragraph()
    _add_college_table(doc, '3.1  Reach Schools',  'C00000', reach_colleges)
    _add_college_table(doc, '3.2  Match Schools',  '375623', match_colleges)
    _add_college_table(doc, '3.3  Safety Schools', '1F3864', safety_colleges)

    _add_section_heading(doc, 4, 'Profile-Building Activities')

    _add_sub_heading(doc, '4.1  Recommended Activities and Competitions')
    for item in activities:
        _add_bullet(doc, item)
    doc.add_paragraph()

    _add_sub_heading(doc, '4.2  Leadership Development')
    for item in leadership_recs:
        _add_bullet(doc, item)
    doc.add_paragraph()

    _add_sub_heading(doc, '4.3  Community Service')
    for item in service_recs:
        _add_bullet(doc, item)
    doc.add_paragraph()

    _add_section_heading(doc, 5, 'Financial Aid and Scholarships')

    _add_sub_heading(doc, '5.1  Federal Aid Overview')
    for item in federal_aid_notes:
        _add_bullet(doc, item)
    doc.add_paragraph()

    _add_sub_heading(doc, '5.2  State Grant Programs')
    for item in state_aid_notes:
        _add_bullet(doc, item)
    doc.add_paragraph()

    _add_sub_heading(doc, '5.3  Institutional Aid Highlights')
    for item in institutional_aid:
        _add_bullet(doc, item)
    doc.add_paragraph()

    _add_sub_heading(doc, '5.4  External Scholarships')
    _add_scholarship_table(doc, scholarships)

    _add_section_heading(doc, 6, 'Application Timeline and Checklist')

    _add_sub_heading(doc, '6.1  Month-by-Month Action Plan')
    _add_timeline_table(doc, timeline_rows)

    _add_sub_heading(doc, '6.2  Application Strategy Recommendations')
    for item in strategy_notes:
        _add_bullet(doc, item)
    doc.add_paragraph()

    out_path = f'output/{student_name.replace(" ", "_")}_plan.docx'
    doc.save(out_path)
    print(f'Plan saved to {out_path}')
    return out_path


# ---------------------------------------------------------------------------
# Prathiksha Tamilmani — College Plan
# ---------------------------------------------------------------------------

if __name__ == '__main__':
    generate_plan(
        student_name   = 'Prathiksha Tamilmani',
        grade          = '11th (Rising)',
        city           = 'Irving',
        state          = 'TX',
        intended_major = 'Business Administration / Healthcare Management / Supply Chain',

        career_paths = [
            dict(
                title        = 'Healthcare Consultant',
                fit_label    = 'Best Fit',
                major        = 'Business Administration, Healthcare Management, or Health Informatics',
                grad_school  = 'Strongly recommended — MBA or MHA (Master of Health Administration) within 5 years',
                salary_start = '$65,000 – $85,000',
                salary_mid   = '$110,000 – $160,000',
                outlook      = 'Growing rapidly — 28% projected growth over next decade (BLS)',
                why_fit      = 'Combines your interest in both business and healthcare. Your DECA background, IB curriculum, and interest in consulting make this a natural fit. You do not need clinical training — this is a business-side role in healthcare.',
                day_in_life  = 'Advise hospitals and health systems on operations, cost reduction, technology, and strategy. Work with a team to analyze data, build presentations, and present recommendations to hospital executives. Frequent client travel once senior.',
            ),
            dict(
                title        = 'Supply Chain Manager (Healthcare or Corporate)',
                fit_label    = 'Strong Alternative',
                major        = 'Supply Chain Management, Business Administration, or Operations Management',
                grad_school  = 'Optional — MBA can accelerate to Director level',
                salary_start = '$60,000 – $75,000',
                salary_mid   = '$95,000 – $130,000',
                outlook      = 'Growing — 28% increase in logistics/supply chain roles (BLS); healthcare supply chain is especially in demand post-pandemic',
                why_fit      = 'Your interest in supply chain directly maps here. DFW is a major logistics hub — Dallas is home to major corporate HQs (Toyota, AT&T, American Airlines) and hospital systems (Baylor Scott & White, Texas Health Resources) all hiring supply chain talent.',
                day_in_life  = 'Manage the flow of goods, services, and information across organizations. In healthcare: ensure hospitals have the right equipment and pharmaceuticals at the right time. In corporate: optimize manufacturing and distribution networks.',
            ),
            dict(
                title        = 'Hospital Administrator / Health Services Manager',
                fit_label    = 'Alternative Path',
                major        = 'Health Administration, Public Health, or Business Administration',
                grad_school  = 'Strongly recommended — MHA or MBA with health concentration',
                salary_start = '$55,000 – $70,000',
                salary_mid   = '$100,000 – $145,000',
                outlook      = 'Growing — 28% over next decade; one of the fastest-growing management fields (BLS)',
                why_fit      = 'Your interest in hospital administration and public health aligns directly here. Texas Medical Center (Houston) and Baylor Scott & White (Dallas) are two of the largest health systems in the country — excellent for internships and career starts.',
                day_in_life  = 'Run the business operations of a hospital, clinic, or public health agency. Manage budgets, staff, patient experience, and compliance. Work with doctors but focus on the organizational side, not clinical care.',
            ),
        ],

        course_recs = [
            'IB Business Management HL (if not already taken) — directly builds your business foundation',
            'IB Economics HL — essential for consulting and healthcare finance',
            'IB Mathematics: Analysis and Approaches HL — strengthens analytical skills valued by top business programs',
            'AP Statistics — strong complement to supply chain and health data analysis',
            'Dual Enrollment: Introduction to Business or Accounting at a local community college (fall/spring)',
            'Consider IB Extended Essay in a business or health systems topic — impressive for top college applications',
        ],

        gpa_targets = [
            'Current GPA: 3.55 unweighted — target 3.7+ unweighted by end of 11th grade',
            'IB Diploma predicted score: aim for 34+ out of 45 (roughly equivalent to 3.7+ at most colleges)',
            'SAT target for match schools: 1200–1300; for reach schools: 1350+',
            'SAT target for UT Austin McCombs School of Business: 1300+ recommended',
            'PSAT score of 960 = SAT equivalent of ~1000; with focused prep, a 1200–1300 by August is achievable',
            'Recommended SAT prep: Khan Academy free SAT prep (links to College Board), Princeton Review, or local prep class in Irving/DFW',
            'Plan to retake SAT in October or November 2025 if August score is below target — most students improve 50–100 points on retake',
        ],

        summer_programs = [
            '--- NATIONAL PROGRAMS ---',
            'Bank of America Student Leaders (Rising 12th): Paid internship + leadership summit — deadline Feb; apply Nov 2025',
            'DECA Leadership Academy: Summer leadership program for DECA members — aligns directly with your club involvement',
            'Wharton Leadership in the Business World (Rising 12th): One of the most prestigious HS business programs; ~$5,000 but financial aid available; apply Feb 2026',
            'NFTE (Network for Teaching Entrepreneurship): Free entrepreneurship program; aligns with your business interest',
            '',
            '--- LOCAL / NEARBY (Irving / DFW) ---',
            'Baylor Scott & White Health (Dallas/Irving) — Teen Volunteer Program: Apply Feb–Mar; call their volunteer services directly; highly regarded on applications for health administration interest',
            'Texas Health Resources (Arlington/DFW area) — Junior Volunteer Program: One of largest TX health systems; teen volunteer spots fill quickly — apply February',
            'UT Dallas (Richardson, TX) — Pre-college business or healthcare summer programs for TX students: UTD is 20 min from Irving; check utdallas.edu for HS programs',
            'TCU (Fort Worth) — Neeley School of Business summer programs: 30 min from Irving; check neeley.tcu.edu/academics/hs-programs',
            'Irving City Government — Summer Intern/Aide: Email city.irving.tx.us — city management, public health, or administration; unpaid but impressive',
            'DFW Airport (Irving) — Supply Chain / Logistics exposure: World\'s 4th busiest airport is headquartered in Irving; contact their HR for student observation programs',
            'Local Chamber of Commerce (Irving): irvingtexas.com/chamber — lists member businesses that may accept high school interns',
            '',
            '--- SELF-DIRECTED (FREE) ---',
            'Complete Google\'s free "Fundamentals of Digital Marketing" certification online (40 hours, internationally recognized)',
            'Take an Introduction to Supply Chain Management course on Coursera (free to audit)',
            'Start a personal project: create a simple business plan for a hypothetical company or analyze a local business — can anchor your college essay',
            'Read: "The Innovators Dilemma" (Clayton Christensen) or "Being Mortal" (Atul Gawande) — shows intellectual depth beyond coursework',
        ],

        reach_colleges = [
            dict(
                name            = 'University of Pennsylvania (Wharton)',
                location        = 'Philadelphia, PA',
                type            = 'Private Ivy League',
                acceptance_rate = '7%',
                avg_gpa         = '3.9',
                sat_act_range   = '1500–1570 / 34–36',
                enrollment      = '10,000 undergrad',
                program_notes   = '#1 undergraduate business program in the US; strong healthcare management track',
                annual_coa      = '$87,000',
                net_price       = '~$25,000–$35,000 (est. at >$150K income)',
                aid_notes       = 'Meets 100% of demonstrated need; merit scholarships limited at Wharton',
                fit_reason      = 'Top business school with healthcare concentration; DECA and IB background will stand out',
            ),
            dict(
                name            = 'Georgetown University (McDonough)',
                location        = 'Washington, DC',
                type            = 'Private Research',
                acceptance_rate = '12%',
                avg_gpa         = '3.9',
                sat_act_range   = '1380–1550 / 32–35',
                enrollment      = '7,500 undergrad',
                program_notes   = 'Strong business + health policy intersection; proximity to federal health agencies (NIH, HHS)',
                annual_coa      = '$83,000',
                net_price       = '~$35,000 (est.)',
                aid_notes       = 'Merit scholarships available; strong first-gen support programs',
                fit_reason      = 'Perfect overlap of business + public health; DC location ideal for health policy internships',
            ),
            dict(
                name            = 'Emory University (Goizueta Business School)',
                location        = 'Atlanta, GA',
                type            = 'Private Research',
                acceptance_rate = '11%',
                avg_gpa         = '3.8',
                sat_act_range   = '1390–1530 / 32–35',
                enrollment      = '7,000 undergrad',
                program_notes   = 'Adjacent to CDC and Emory Healthcare — unmatched for health + business; strong supply chain program',
                annual_coa      = '$79,000',
                net_price       = '~$30,000 (est.)',
                aid_notes       = 'Emory Scholarship — up to full tuition for top merit candidates',
                fit_reason      = 'Healthcare + business + supply chain all in one; CDC next door is a pipeline for public health internships',
            ),
            dict(
                name            = 'Rice University (Jones School)',
                location        = 'Houston, TX',
                type            = 'Private Research',
                acceptance_rate = '9%',
                avg_gpa         = '3.96',
                sat_act_range   = '1500–1580 / 34–36',
                enrollment      = '4,200 undergrad',
                program_notes   = 'Adjacent to Texas Medical Center (largest medical complex in the world); strong business + health policy',
                annual_coa      = '$72,000',
                net_price       = '~$20,000–$30,000 (est.)',
                aid_notes       = 'Meets 100% of demonstrated need; strong merit scholarships',
                fit_reason      = 'Texas school (in-state culture); Texas Medical Center internship access is unmatched nationally for health administration',
            ),
        ],

        match_colleges = [
            dict(
                name            = 'University of Texas at Austin (McCombs)',
                location        = 'Austin, TX',
                type            = 'Public — Texas flagship',
                acceptance_rate = '31%',
                avg_gpa         = '3.75',
                sat_act_range   = '1230–1490 / 28–34',
                enrollment      = '51,000',
                program_notes   = '#5 undergraduate business program (US News); strong supply chain and health management tracks',
                annual_coa      = '$32,000 (in-state)',
                net_price       = '~$20,000–$25,000 (est. at >$150K)',
                aid_notes       ='Numerous merit scholarships; Forty Acres Scholars (full ride, very competitive)',
                fit_reason      = 'Best value in Texas for business; IB diploma is a strong differentiator at UT; DFW applicants well-represented',
            ),
            dict(
                name            = 'Texas A&M University (Mays Business School)',
                location        = 'College Station, TX',
                type            = 'Public — Texas flagship',
                acceptance_rate = '63%',
                avg_gpa         = '3.65',
                sat_act_range   = '1160–1380 / 26–32',
                enrollment      = '74,000',
                program_notes   = 'Top-ranked supply chain management program in the US; strong DECA alumni network',
                annual_coa      = '$30,000 (in-state)',
                net_price       = '~$18,000–$22,000 (est.)',
                aid_notes       = 'Academic Excellence Award; President\'s Achievement Award for merit',
                fit_reason      = 'Best supply chain program in Texas; DECA experience and IB curriculum make you a competitive applicant',
            ),
            dict(
                name            = 'Southern Methodist University (Cox School)',
                location        = 'Dallas, TX',
                type            = 'Private',
                acceptance_rate = '52%',
                avg_gpa         = '3.7',
                sat_act_range   = '1250–1450 / 28–33',
                enrollment      = '12,000',
                program_notes   = 'Located in Dallas; strong healthcare management and consulting tracks; excellent DFW employer network',
                annual_coa      = '$79,000',
                net_price       = '~$35,000–$45,000 (est.)',
                aid_notes       = 'President\'s Scholars — full tuition; Dean\'s Scholars — half tuition; aggressive merit aid',
                fit_reason      = 'In Dallas — you stay close to home; SMU Cox has one of the best DFW corporate recruiting networks; strong merit scholarships',
            ),
            dict(
                name            = 'University of Texas at Dallas (Naveen Jindal School)',
                location        = 'Richardson, TX',
                type            = 'Public',
                acceptance_rate = '79%',
                avg_gpa         = '3.6',
                sat_act_range   = '1170–1410 / 26–32',
                enrollment      = '29,000',
                program_notes   = '#1 supply chain program in Texas (some rankings); strong healthcare management; 20 min from Irving',
                annual_coa      = '$28,000 (in-state)',
                net_price       = '~$15,000–$20,000 (est.)',
                aid_notes       = 'McDermott Scholars (full ride, very competitive); Academic Excellence Scholarship',
                fit_reason      = 'Closest top business school to Irving; supply chain program is nationally ranked; excellent for staying in DFW',
            ),
            dict(
                name            = 'University of Houston (Bauer College)',
                location        = 'Houston, TX',
                type            = 'Public',
                acceptance_rate = '65%',
                avg_gpa         = '3.5',
                sat_act_range   = '1100–1320 / 23–29',
                enrollment      = '46,000',
                program_notes   = 'AACSB accredited; strong supply chain and healthcare administration tracks; proximity to Texas Medical Center',
                annual_coa      = '$26,000 (in-state)',
                net_price       = '~$14,000–$18,000 (est.)',
                aid_notes       = 'Tier One Scholarship; numerous merit awards for strong IB/AP students',
                fit_reason      = 'Texas Medical Center (next door) provides unbeatable internship access for health administration interest',
            ),
            dict(
                name            = 'Texas Christian University (Neeley School)',
                location        = 'Fort Worth, TX',
                type            = 'Private',
                acceptance_rate = '44%',
                avg_gpa         = '3.6',
                sat_act_range   = '1180–1390 / 26–31',
                enrollment      = '11,000',
                program_notes   = 'Strong supply chain and entrepreneurship; DFW location; AACSB accredited',
                annual_coa      = '$72,000',
                net_price       = '~$30,000–$40,000 (est.)',
                aid_notes       = 'Chancellor\'s Scholarship (full tuition); Dean\'s Scholarship (half tuition)',
                fit_reason      = '30 minutes from Irving; strong merit scholarships and DFW employer network for business careers',
            ),
        ],

        safety_colleges = [
            dict(
                name            = 'University of North Texas (G. Brint Ryan College)',
                location        = 'Denton, TX',
                type            = 'Public',
                acceptance_rate = '75%',
                avg_gpa         = '3.3',
                sat_act_range   = '1010–1230 / 21–27',
                enrollment      = '44,000',
                program_notes   = 'Strong supply chain and health services management programs; 30 min from Irving',
                annual_coa      = '$24,000 (in-state)',
                net_price       = '~$12,000–$16,000 (est.)',
                aid_notes       = 'Honors College scholarships; competitive merit awards for IB students',
                fit_reason      = 'Close to home; solid business programs; your 3.55 GPA and IB curriculum make you a strong candidate for Honors College',
            ),
            dict(
                name            = 'Texas Woman\'s University',
                location        = 'Denton, TX',
                type            = 'Public',
                acceptance_rate = '85%',
                avg_gpa         = '3.3',
                sat_act_range   = '980–1190 / 19–26',
                enrollment      = '16,000',
                program_notes   = 'Strong health administration and public health programs; note: open to all genders despite name',
                annual_coa      = '$22,000 (in-state)',
                net_price       = '~$10,000–$14,000 (est.)',
                aid_notes       = 'Academic merit scholarships; Pioneer Women\'s Leadership scholarships',
                fit_reason      = 'Strong health administration specialty; very close to Irving; affordable with merit aid',
            ),
            dict(
                name            = 'University of Texas at Arlington (College of Business)',
                location        = 'Arlington, TX',
                type            = 'Public',
                acceptance_rate = '72%',
                avg_gpa         = '3.3',
                sat_act_range   = '990–1220 / 20–27',
                enrollment      = '43,000',
                program_notes   = 'Strong supply chain and healthcare management; 20 min from Irving; excellent DFW employer connections',
                annual_coa      = '$23,000 (in-state)',
                net_price       = '~$11,000–$15,000 (est.)',
                aid_notes       = 'Maverick Excellence Scholarship; Honors College scholarship',
                fit_reason      = 'Closest university to Irving; strong business programs; guaranteed safety with your profile',
            ),
            dict(
                name            = 'Texas State University (McCoy College)',
                location        = 'San Marcos, TX',
                type            = 'Public',
                acceptance_rate = '88%',
                avg_gpa         = '3.25',
                sat_act_range   = '980–1190 / 19–26',
                enrollment      = '38,000',
                program_notes   = 'AACSB accredited; strong healthcare administration and supply chain programs',
                annual_coa      = '$25,000 (in-state)',
                net_price       = '~$12,000–$16,000 (est.)',
                aid_notes       = 'Presidential Scholarship (full in-state tuition); Academic Excellence Award',
                fit_reason      = 'Strong safety option; AACSB accredited business school; good merit scholarship opportunities for your profile',
            ),
        ],

        activities = [
            'DECA — Compete at state and national level (DECA ICDC); aim for officer or VP role in 11th grade',
            'Business or Entrepreneurship Club — if not at your school, consider starting one (founding a club = leadership signal to colleges)',
            'HOSA (Health Occupations Students of America) — directly aligns with healthcare + business interest; check if your school has a chapter',
            'Model United Nations — develops the policy and public health thinking relevant to your healthcare interest',
            'Swimming — continue competitively if possible; or transition to managing/leading the team (captain, coach assistant) for leadership signal',
            'UIL Accounting or Business competitions — available across Texas; directly tested on Common App',
            'Junior Achievement Company Program — run a mini-company for a semester; aligns with DECA and entrepreneurship interest',
            'National Business Honor Society — check eligibility (typically 3.0+ GPA in business courses)',
        ],

        leadership_recs = [
            'Run for DECA Chapter President or VP before senior year — leadership title carries significant weight',
            'Volunteer as a youth swim coach assistant or teach swim lessons — converts your hobby into leadership + community service',
            'Apply to Irving ISD or local city government Youth Advisory Board — direct public policy and administration experience',
            'Start or lead a first-generation college student awareness club at your school — shows initiative and ties to your identity as a first-gen student',
        ],

        service_recs = [
            'Continue food bank volunteering — try to increase hours and take on a coordinator or logistics role (aligns with supply chain interest)',
            'Organize a food or supply drive at your school — you manage the supply chain, promote it, and deliver results (DECA + supply chain in action)',
            'Volunteer at a local free clinic or community health fair in Irving/DFW — directly relevant to healthcare administration interest',
            'Irving Cares (local nonprofit in Irving) — provides social services to Irving residents; volunteer or intern in their operations',
            'Aim for 150–200 total community service hours by end of 11th grade — strong for Texas scholarship applications',
        ],

        federal_aid_notes = [
            'FAFSA (Free Application for Federal Student Aid): File on October 1, 2026 for fall 2027 enrollment — do not wait',
            'At household income >$150K, you are unlikely to qualify for Pell Grants (need-based federal aid)',
            'You may still qualify for Federal Direct Unsubsidized Loans (not need-based) — maximum $5,500/year freshman year',
            'Work-Study eligibility is need-based — likely not available at your income level',
            'CSS Profile: Required by many private colleges (Georgetown, Rice, Emory, Penn) in addition to FAFSA — file this simultaneously in October',
            'Focus your financial aid strategy on merit scholarships, which are not income-based',
        ],

        state_aid_notes = [
            'Texas TEXAS Grant: Need-based; unlikely to qualify at >$150K household income — but file FAFSA anyway as income verification rules vary',
            'Texas B-On-Time Loan: Need-based loan with forgiveness provision — check eligibility when attending TX public college',
            'Texas Public Education Grant (TPEG): Administered by individual universities; some have merit components — check each school',
            'UT Austin Longhorn Scholarship: Merit-based; competitive; up to $5,000/year for academically strong TX students',
            'Texas A&M Academic Excellence Award: Merit-based; up to $3,000/year; apply through admissions process',
            'Note: For Texas public universities, in-state tuition (~$12,000–$15,000/year vs ~$38,000 out-of-state) is itself a major financial benefit',
        ],

        institutional_aid = [
            'SMU Cox School — President\'s Scholars: Full tuition (~$65,000/year); very competitive; requires 1400+ SAT and 3.9+ GPA',
            'SMU Cox School — Dean\'s Scholars: Half tuition; requires 1300+ SAT and strong profile — within your reach with strong SAT prep',
            'TCU Neeley — Chancellor\'s Scholarship: Full tuition; very competitive; strong IB background helps',
            'TCU Neeley — Dean\'s Scholarship: Half tuition; your DECA + IB profile is competitive',
            'UT Dallas — McDermott Scholars: Full tuition + stipend; extremely competitive; apply November; strong IB/AP students preferred',
            'UT Dallas — Academic Excellence Scholarship: $5,000–$10,000/year; merit-based; competitive for IB students',
            'Rice University: Meets 100% of demonstrated need; also offers merit scholarships — unusual combination for a top-10 school',
            'Emory University — Emory Scholarship: Up to full tuition; competitive; IB Diploma applicants viewed favorably',
        ],

        scholarships = [
            dict(
                name        = 'National Merit Scholarship',
                amount      = 'Up to $2,500 + corporate/college sponsors (some full rides)',
                eligibility = 'Top PSAT scorers — TX cutoff typically ~220 Selection Index; file PSAT Oct 2025 (11th grade)',
                deadline    = 'PSAT taken October 2025 — automatic entry',
                website     = 'nationalmerit.org',
            ),
            dict(
                name        = 'Coca-Cola Scholars Program',
                amount      = '$20,000',
                eligibility = 'High achieving seniors; community leadership; 3.0+ GPA; US citizen',
                deadline    = 'October 31 (senior year)',
                website     = 'coca-colascholarsfoundation.org',
            ),
            dict(
                name        = 'Jack Kent Cooke Foundation College Scholarship',
                amount      = 'Up to $55,000/year',
                eligibility = 'High-achieving students with financial need; first-gen students prioritized; GPA 3.5+',
                deadline    = 'November (senior year)',
                website     = 'jkcf.org',
            ),
            dict(
                name        = 'DECA Scholarships',
                amount      = '$1,000 – $25,000 (varies by sponsor)',
                eligibility = 'Active DECA members; academic achievement; DECA competition performance',
                deadline    = 'February–March (senior year)',
                website     = 'deca.org/scholarships',
            ),
            dict(
                name        = 'Hispanic Scholarship Fund',
                amount      = '$500 – $5,000',
                eligibility = 'Hispanic heritage; 3.0+ GPA; US citizen or permanent resident; financial need considered',
                deadline    = 'February 15 (senior year)',
                website     = 'hsf.net',
            ),
            dict(
                name        = 'QuestBridge College Prep Scholars',
                amount      = 'Recognition + access to QuestBridge National College Match (full ride potential)',
                eligibility = 'High-achieving, lower-income juniors; however, income >$150K likely disqualifies need component',
                deadline    = 'March (junior year — apply now)',
                website     = 'questbridge.org',
            ),
            dict(
                name        = 'Bank of America Preferred Scholars',
                amount      = '$3,500/year renewable',
                eligibility = 'Enrolled at select universities; merit-based; no income requirement',
                deadline    = 'Varies by school — check with enrolled university',
                website     = 'bankofamerica.com/about/community-giving',
            ),
            dict(
                name        = 'Toyota Scholarship (Irving, TX)',
                amount      = 'Varies — typically $2,500–$5,000',
                eligibility = 'Texas students; Toyota is headquartered in Plano, TX (near Irving); check Toyota USA Foundation programs',
                deadline    = 'Spring of senior year',
                website     = 'toyotausafoundation.org',
            ),
            dict(
                name        = 'Texas Business Hall of Fame Scholarship',
                amount      = '$7,500 – $10,000',
                eligibility = 'Texas resident; pursuing business degree at TX university; community involvement; 3.0+ GPA',
                deadline    = 'January (senior year)',
                website     = 'tbhf.com/scholarships',
            ),
            dict(
                name        = 'IB Diploma Merit Awards (various universities)',
                amount      = '$2,000 – $10,000/year',
                eligibility = 'IB Diploma earners; varies by university — UT Austin, SMU, TCU, and UTD all offer IB-specific merit awards',
                deadline    = 'Automatic with admissions application — note IB on Common App',
                website     = 'Check each university\'s admissions page for IB merit awards',
            ),
        ],

        timeline_rows = [
            ('Summer 2025 (Now)',
             'Take SAT prep classes in Irving/DFW; write first SAT in August 2025; complete dual enrollment Federal Government course; volunteer at local hospital or food bank; explore HOSA membership for fall'),
            ('Aug–Sep 2025 (11th Grade Starts)',
             'Join or run for DECA officer; explore HOSA; request to shadow a hospital administrator in DFW; research college list; begin college visits to UTD, SMU, TCU (all nearby)'),
            ('Oct 2025',
             'Take PSAT (National Merit qualifying exam — critical for your junior year); retake SAT if August score was below 1200; review PSAT score for National Merit trajectory'),
            ('Nov–Dec 2025',
             'Research summer programs for 2026 (Bank of America Student Leaders, Wharton LBW — deadlines are Feb); apply to QuestBridge College Prep Scholars (March deadline); research IB Extended Essay topic'),
            ('Jan–Feb 2026',
             'Apply to Bank of America Student Leaders (Feb deadline); apply to Wharton LBW or equivalent business summer program; submit any spring scholarship applications; plan campus visits to TX match/safety schools'),
            ('Mar–May 2026',
             'Finalize college list (reach/match/safety); take AP exams; complete IB assessments; plan senior year course load; take SAT one more time if needed'),
            ('Summer 2026 (Rising 12th)',
             'Complete prestigious internship or summer program (this is the most important summer); draft Common App personal statement (aim for draft 1 by Aug 1); request letters of recommendation from 2 teachers before summer ends'),
            ('Aug 1, 2026',
             'Common App opens — fill out activities section immediately; finalize personal statement; begin Why Us essays for each school'),
            ('Sep–Oct 2026',
             'Submit Early Action applications (due Nov 1–15); file FAFSA on Oct 1; file CSS Profile for private schools; finalize scholarship applications'),
            ('Nov 1–15, 2026',
             'EA application deadlines — submit to all EA schools; confirm all scholarship applications submitted'),
            ('Dec 2026',
             'EA decisions arrive; complete remaining Regular Decision applications; compare any early financial aid estimates'),
            ('Jan 1–15, 2027',
             'RD application deadlines — submit all remaining applications; apply for Texas-specific scholarships (TBHF, Toyota, etc.)'),
            ('Feb–Mar 2027',
             'Financial aid award letters arrive; compare net price across schools; negotiate aid if needed (call financial aid offices — this works)'),
            ('Apr 2027',
             'Campus visits to top choices; compare final offers; decide'),
            ('May 1, 2027',
             'National Decision Day — submit enrollment deposit to chosen school'),
        ],

        strategy_notes = [
            'Apply Early Action (non-binding) to all schools where possible — acceptance rates are typically 10–15% higher than Regular Decision, and you will hear back by December.',
            'Your IB Diploma is a meaningful differentiator — explicitly highlight it in the activities section and supplemental essays. Many admissions officers consider IB more rigorous than AP.',
            'As a first-generation college student, note this clearly on every application. Many schools have dedicated admissions staff and scholarships specifically for first-gen students.',
            'The DECA club is your strongest current activity — go deep, not wide. Competing at state or national DECA level and holding an officer title is far more impressive than joining five clubs.',
            'For Texas public schools (UT Austin, Texas A&M, UTD, UTA, UNT), automatic admission policies may apply — UT Austin admits top 6% of TX high school class automatically; check if your school\'s rank qualifies.',
            'Target SAT score of 1300+ to be competitive for SMU Dean\'s Scholarship and UT Austin McCombs direct admission — this is your most important academic goal for summer 2025.',
            'Do not skip the "Additional Information" section on Common App — use it to explain your dual enrollment, IB extended essay, and any other academic context that does not fit elsewhere.',
            'For the personal statement, consider writing about the intersection of business and healthcare — your interest in systems and how organizations work. Specific, concrete stories beat general statements.',
        ],
    )
